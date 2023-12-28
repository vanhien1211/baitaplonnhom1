import tkinter as tk
from tkinter import filedialog
import sympy as sym
from tkinter import messagebox
from docx import Document
from tkinter import scrolledtext
import re

global_word_content = ""
result_text = ""

def open_file_dialog():
    global global_word_content
    file_path = filedialog.askopenfilename(filetypes=[("Word Files", "*.docx")])
    if file_path:
        read_word_file(file_path)
        text_widget.delete("1.0", "end")
        for line in global_word_content:
            text_widget.insert(tk.END, line + "\n")

def read_word_file(file_path):
    global global_word_content
    doc = Document(file_path)
    global_word_content = [paragraph.text for paragraph in doc.paragraphs]

def display_result(result, display):
    display.delete("1.0", "end-1c")
    for i in result:
        display.insert(tk.END, i)

def lookup_sympify(display):
    chuoi_ham = text_widget.get("1.0", "end-1c")
    if chuoi_ham=='':
        return -1
    else:
        chuoi_moi = re.sub(r'(Bài|Câu|bài|câu) \d{1,10}:', '', chuoi_ham)
        hams = re.split(r'[\n\s,;]+', chuoi_moi)
        hams = [item for item in hams if item != '']
        return hams

def diff():
    try:
        global result_text
        result_text = []
        hams=lookup_sympify(text_widget)
        if hams == -1:
            messagebox.showwarning("Chú ý", "Hãy nhập biểu thức với các chữ số và biến x")
        else:
            for ham in hams:
                if all(char.isdigit() or char == 'x' or char in ['+', '-', '*', '/', '^', '(', ')', '.'] for char in ham):
                    x = sym.Symbol('x')
                    fx = sym.sympify(ham)
                    daoham1 = sym.diff(fx, x)
                    if isinstance(daoham1, (sym.core.numbers.Number, sym.core.numbers.Float, sym.core.numbers.Integer)):
                        daoham = round(daoham1,2)
                    else:
                        daoham = daoham1
                    result_text.append(f"Đạo hàm của biểu thức {ham} là: {daoham}\n")
                else:
                    ham_loi=ham
                    messagebox.showwarning("Chú ý", f"Biểu thức bị sai {ham_loi}\r\nVui lòng nhập lại hàm này")
            display_result(result_text, result_widget)
    except ValueError:
        messagebox.showwarning("Chú ý", f"Biểu thức bị sai\r\nVui lòng nhập lại hàm này")

def integrate():
    try:
        global result_text
        result_text = []
        hams = lookup_sympify(text_widget)
        if hams == -1:
            messagebox.showwarning("Chú ý", "Hãy nhập biểu thức với các chữ số và biến x")
        else:
            for ham in hams:
                if all(char.isdigit() or char == 'x' or char in ['+', '-', '*', '/', '^', '(', ')', '.'] for char in ham):
                    x = sym.Symbol('x')
                    fx = sym.sympify(ham)
                    tichphan1 = sym.integrate(fx, x)
                    if isinstance(tichphan1, (sym.core.numbers.Number, sym.core.numbers.Float, sym.core.numbers.Integer)):
                        tichphan = round(tichphan1, 2)
                    else:
                        tichphan = tichphan1
                    result_text.append(f"Tích phân của biểu thức {ham} là: {tichphan}\n")
                else:
                    ham_loi=ham
                    messagebox.showwarning("Chú ý", f"Biểu thức bị sai {ham_loi}\r\nVui lòng nhập lại hàm này")
            display_result(result_text, result_widget)
    except ValueError:
        messagebox.showwarning("Chú ý", f"Biểu thức bị sai\r\nVui lòng nhập lại hàm này")

def limit():
    try:
        global result_text
        result_text = []
        hams = lookup_sympify(text_widget)
        if hams == -1:
            messagebox.showwarning("Chú ý", "Hãy nhập biểu thức với các chữ số và biến x")
        else:
            for ham in hams:
                if all(char.isdigit() or char == 'x' or char in ['+', '-', '*', '/', '^', '(', ')', '.'] for char in ham):
                    x = sym.Symbol('x')
                    fx = sym.sympify(ham)
                    ghan1 = sym.limit(fx, x, 0)
                    if isinstance(ghan1, (sym.core.numbers.Number, sym.core.numbers.Float, sym.core.numbers.Integer)):
                        ghan = round(ghan1, 2)
                    else:
                        ghan = ghan1
                    result_text.append(f"Giới hạn của biểu thức {ham} khi x tiến tới 0 là: {ghan}\n")
                else:
                    ham_loi = ham
                    messagebox.showwarning("Chú ý", f"Biểu thức bị sai {ham_loi}\r\nVui lòng nhập lại hàm này")
            display_result(result_text, result_widget)
    except ValueError:
             messagebox.showwarning("Chú ý", f"Biểu thức bị sai\r\nVui lòng nhập lại hàm này")


def simplify():
    try:
        global result_text
        result_text = []
        hams = lookup_sympify(text_widget)
        if hams == -1:
            messagebox.showwarning("Chú ý", "Hãy nhập biểu thức với các chữ số và biến x")
        else:
            for ham in hams:
                if all(char.isdigit() or char == 'x' or char in ['+', '-', '*', '/', '^', '(', ')', '.'] for char in ham):
                    x = sym.Symbol('x')
                    fx = sym.sympify(ham)
                    rutgon1 = sym.simplify(fx)
                    if isinstance(rutgon1, (sym.core.numbers.Number, sym.core.numbers.Float, sym.core.numbers.Integer)):
                        rutgon = round(rutgon1, 2)
                    else:
                        rutgon = rutgon1
                    result_text.append(f"Rút gọn của biểu thức {ham} là: {rutgon}\n")
                else:
                    ham_loi=ham
                    messagebox.showwarning("Chú ý", f"Biểu thức bị sai {ham_loi}\r\nVui lòng nhập lại hàm này")
            display_result(result_text, result_widget)
    except ValueError:
        messagebox.showwarning("Chú ý", f"Biểu thức bị sai\r\nVui lòng nhập lại hàm này")

def save_to_word_file():
    global result_text
    if result_text == '':
        messagebox.showinfo("Cảnh báo", "Chưa thực hiện phép toán nào")
    else:
        file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Files", "*.docx")])
        if file_path:
            doc = Document()
            doc.add_paragraph(result_text)
            doc.save(file_path)
            print(f"Đã lưu vào: {file_path}")

window = tk.Tk()
window.title("Ứng dụng hỗ trợ môn giải tích")
window.geometry('390x500')
window.resizable(width=False, height=False)

window.columnconfigure(0, weight=1)
window.columnconfigure(1, weight=1)
window.columnconfigure(2, weight=1)


tdaoham = tk.Button(window, text="Tính đạo hàm", command=diff)
tdaoham.grid(row = 0, column =0, padx=5, pady=5)

ttichphan = tk.Button(window, text="Tính tích phân", command=integrate)
ttichphan.grid(row = 0, column =1, padx=5, pady=5)

tghan = tk.Button(window, text="Tính giới hạn", command=limit)
tghan.grid(row = 1, column =0, padx=5, pady=5)

trutgon = tk.Button(window, text="Rút gọn biểu thức", command=simplify)
trutgon.grid(row = 1, column =1, padx=5, pady=5)

button = tk.Button(window, text="Chọn tệp Word", command=open_file_dialog)
button.grid(row = 0, column =2, padx=5, pady=5)

button2 = tk.Button(window, text="Lưu kết quả", command=save_to_word_file)
button2.grid(row = 1, column =2, padx=5, pady=5)

label_word = tk.Label(window, text="Dữ liệu bài toán :")
label_word.grid(row=2, column=0, columnspan=2)

text_widget = scrolledtext.ScrolledText(window, width=40, height=10)
text_widget.grid(row=3, column=0, columnspan=5)

label_result = tk.Label(window, text="Kết quả bài toán :")
label_result.grid(row=4, column=0, columnspan=2)

result_widget = scrolledtext.ScrolledText(window, width=40, height=10)
result_widget.grid(row=5, column=0, columnspan=5)

window.mainloop()
