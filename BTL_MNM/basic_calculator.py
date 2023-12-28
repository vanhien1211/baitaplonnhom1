import tkinter as tk
from tkinter import messagebox, scrolledtext
from tkinter import filedialog
from docx import Document


def open_word_file():
    file_path = filedialog.askopenfilename(filetypes=[("Word Files", "*.docx"), ("All Files", "*.*")])
    if file_path:
        if not file_path.endswith(".docx"):
            messagebox.showinfo("Cảnh báo", "Vui lòng chỉ chọn tệp Word (định dạng .docx).")
            return
        text_widget.delete("1.0", "end-1c")
        doc = Document(file_path)
        content = [paragraph.text for paragraph in doc.paragraphs]
        for line in content:
            text_widget.insert(tk.END, line + "\n")
        noidung = text_widget.get("1.0", "end-1c")
        noidung = noidung[:-1]
        text_widget.delete("1.0", "end-1c")
        text_widget.insert(tk.END, noidung)
        return content
    else:
        messagebox.showinfo("Cảnh báo", "Vui lòng chọn một tệp Word.")


def process_exercises(content):
    exercises = []
    current_exercise = []
    for line in content.splitlines():
        if line.startswith("Bài "):
            if current_exercise:
                exercises.append(current_exercise)
            current_exercise = []
        else:
            current_exercise.append(line)
    if current_exercise:
        exercises.append(current_exercise)
    results = []
    for exercise in exercises:
        results.append(process_word(exercise))
    return results

def process_word(exercise):
    results = []

    for line in exercise:
        if "+" in line:
            operands = line.split("+")
            if len(operands) != 2:
                result = "chỉ thực hiện phép tính 2 số"
            else:
                if operands[0] == '':
                    num1 = 0
                else:
                    num1 = float(operands[0].strip())
                if operands[1] == '':
                    messagebox.showinfo("Cảnh báo", "Nhập lại phép tính")
                    result = None
                else:
                    num2 = float(operands[1].strip())
                    result = num1 + num2
        elif "-" in line:
            operands = line.split("-")
            if len(operands) != 2:
                result = "chỉ thực hiện phép tính 2 số"
            else:
                if operands[0] == '':
                    num1 = 0
                else:
                    num1 = float(operands[0].strip())
                if operands[1] == '':
                    messagebox.showinfo("Cảnh báo", "Nhập lại phép tính")
                    result = None
                else:
                    num2 = float(operands[1].strip())
                    result = num1 - num2
        elif "*" in line:
            operands = line.split("*")
            if len(operands) != 2:
                result = "chỉ thực hiện phép tính 2 số"
            else:
                if operands[0] == '':
                    num1 = 0
                else:
                    num1 = float(operands[0].strip())
                if operands[1] == '':
                    messagebox.showinfo("Cảnh báo", "Nhập lại phép tính")
                    result = None
                else:
                    num2 = float(operands[1].strip())
                    result = num1 * num2
        elif "/" in line:
            operands = line.split("/")
            if len(operands) != 2:
                result = "chỉ thực hiện phép tính 2 số"
            else:
                if operands[0] == '':
                    num1 = 0
                else:
                    num1 = float(operands[0].strip())
                if operands[1] == '':
                    messagebox.showinfo("Cảnh báo", "Nhập lại phép tính")
                    result = None
                else:
                    num2 = float(operands[1].strip())
                    if num2 == 0:
                        result = "không chia được cho 0"
                    else:
                        result = num1 / num2
        else:
            if line.strip().isdigit():
                result = "chỉ thực hiện phép tính 2 số"
            elif line.strip().isalpha():
                result = "chỉ thực hiện phép tính 2 số"
            else:
                messagebox.showinfo("Cảnh báo", "Chỉ nhập cộng, trừ, nhân, chia 2 số")
                return
        results.append(f"{line} = {result}")
    return results

def calculate_and_display_results(text_widget, result_widget):
    if text_widget.get("1.0", "end-1c") == "":
        messagebox.showinfo("Cảnh báo", "Không có dữ liệu từ file Word.")
        return
    content = text_widget.get("1.0", "end")
    results = process_exercises(content)
    result_widget.delete("1.0", "end-1c")
    doc = Document()
    if content.find("Bài")!=-1:
        for i, exercise_results in enumerate(results):
            paragraph = doc.add_paragraph(f"Bài {i + 1}: \n")
            for result in exercise_results:
                paragraph.add_run(result + "\n")
    else:
        paragraph = doc.add_paragraph()
        for result in results:
            paragraph.add_run('\n'.join(result))
    content_after_process = [paragraph.text for paragraph in doc.paragraphs]
    for line in content_after_process:
        result_widget.insert(tk.END, line + "\n")
    return doc

def save_word_file():
    if text_widget.get("1.0", "end-1c") == "":
        messagebox.showinfo("Cảnh báo", "Không có dữ liệu từ file Word.")
        return
    file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Files", "*.docx")])
    if file_path:
        if not file_path.endswith(".docx"):
            messagebox.showinfo("Cảnh báo", "Vui lòng chỉ lưu vào tệp Word (định dạng .docx).")
            return
        doc = calculate_and_display_results(text_widget, result_widget)
        doc.save(file_path)

window = tk.Tk()
window.title("Máy tính cơ bản")
window.geometry('390x500')
window.resizable(width=False, height=False)
window.configure(bg="pink")
window.columnconfigure(0, weight=1)
window.columnconfigure(1, weight=1)
window.columnconfigure(2, weight=1)

button_open_file = tk.Button(window, text="Chọn tệp Word", command=open_word_file,bg='lightblue')
button_open_file.grid(row=0, column=0, padx=10, pady=10)

button_calculate_exercises = tk.Button(window, text="Lưu kết quả", command=save_word_file,bg='lightblue')
button_calculate_exercises.grid(row=0, column=2, padx=10, pady=10)
button_calculate = tk.Button(window, text="Tính toán", command=lambda: calculate_and_display_results(text_widget, result_widget),bg='lightblue')
button_calculate.grid(row=0, column=1, padx=10,pady=10)

label_word = tk.Label(window, text="Dữ liệu đầu vào( Phép tính 2 số): ",bg='pink')
label_word.grid(row=1, column=0, columnspan=2, padx=5, pady=5)

text_widget = scrolledtext.ScrolledText(window, width=40, height=10)
text_widget.grid(row=2, column=0, columnspan=3, padx=5, pady=5)

label_word = tk.Label(window, text="Dữ liệu sau tính toán:",bg='pink')
label_word.grid(row=3, column=0, columnspan=2, padx=5, pady=5)

result_widget = scrolledtext.ScrolledText(window, width=40, height=10)
result_widget.grid(row=4, column=0, columnspan=3, padx=5, pady=5)

window.mainloop()
